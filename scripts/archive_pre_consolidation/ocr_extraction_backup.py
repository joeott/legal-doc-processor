# ocr_extraction.py
import os
import logging
# PIL Image removed - not needed for Stage 1 cloud-only deployment
import PyPDF2 # For page count
import fitz # PyMuPDF
import torch
import time
import boto3
import uuid
import hashlib
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional, Tuple
from botocore.exceptions import ClientError

from scripts.config import (
    QWEN2_VL_OCR_PROMPT, QWEN2_VL_OCR_MAX_NEW_TOKENS,  # Qwen2-VL-OCR configs
    DEPLOYMENT_STAGE, OPENAI_API_KEY, USE_OPENAI_FOR_AUDIO_TRANSCRIPTION,  # Stage management
    STAGE_CLOUD_ONLY, S3_PRIMARY_DOCUMENT_BUCKET, TEXTRACT_USE_ASYNC_FOR_PDF,
    TEXTRACT_CONFIDENCE_THRESHOLD, AWS_DEFAULT_REGION, REDIS_OCR_CACHE_TTL,
    IMAGE_EXTENSIONS, DOCUMENT_EXTENSIONS, AUDIO_EXTENSIONS, VIDEO_EXTENSIONS  # File type detection
)
from scripts.models_init import (
    get_qwen2_vl_ocr_model, get_qwen2_vl_ocr_processor, 
    get_qwen2_vl_ocr_device, get_process_vision_info,
    get_whisper_model, should_load_local_models
) # Use accessor functions
from scripts.supabase_utils import generate_document_url, SupabaseManager
from scripts.textract_utils import TextractProcessor
from scripts.s3_storage import S3StorageManager
from scripts.redis_utils import redis_cache, get_redis_manager
import tempfile
import requests

# OpenAI import for Stage 1
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

logger = logging.getLogger(__name__)


def detect_file_category(file_path: str) -> str:
    """
    Categorize file into image, document, audio, video, or unknown based on extension.
    
    Args:
        file_path: Path to the file or filename
        
    Returns:
        str: 'image', 'document', 'audio', 'video', or 'unknown'
    """
    try:
        from pathlib import Path
        extension = Path(file_path).suffix.lower()
        
        if extension in IMAGE_EXTENSIONS:
            return 'image'
        elif extension in DOCUMENT_EXTENSIONS:
            return 'document'
        elif extension in AUDIO_EXTENSIONS:
            return 'audio'
        elif extension in VIDEO_EXTENSIONS:
            return 'video'
        else:
            logger.warning(f"Unknown file type for extension '{extension}' in file: {file_path}")
            return 'unknown'
    except Exception as e:
        logger.error(f"Error detecting file category for {file_path}: {e}")
        return 'unknown'


def is_image_file(file_path: str) -> bool:
    """
    Check if a file is an image based on its extension.
    
    Args:
        file_path: Path to the file or filename
        
    Returns:
        bool: True if file is an image, False otherwise
    """
    return detect_file_category(file_path) == 'image'


def get_supported_file_types() -> Dict[str, List[str]]:
    """
    Get all supported file types organized by category.
    
    Returns:
        dict: Dictionary with categories as keys and lists of extensions as values
    """
    return {
        'image': list(IMAGE_EXTENSIONS),
        'document': list(DOCUMENT_EXTENSIONS),
        'audio': list(AUDIO_EXTENSIONS),
        'video': list(VIDEO_EXTENSIONS)
    }


def render_pdf_page_to_image(pdf_path: str, page_number: int):
    # This function is not used in Stage 1 cloud-only deployment
    # PIL dependency has been removed
    logger.warning("render_pdf_page_to_image called but PIL not available in Stage 1")
    return None

def extract_text_from_pdf_qwen_vl_ocr(pdf_path: str) -> tuple[str | None, list | None]:
    """
    Extract text from PDF using Qwen2-VL-OCR model.
    In Stage 1, this function is bypassed in favor of Mistral OCR.
    """
    # Stage 1: Skip local model processing
    if DEPLOYMENT_STAGE == "1" or not should_load_local_models():
        logger.info("Stage 1 deployment detected. Skipping Qwen2-VL-OCR in favor of cloud OCR services.")
        return None, None
    
    # Get models using accessor functions
    QWEN2_VL_OCR_MODEL = get_qwen2_vl_ocr_model()
    QWEN2_VL_OCR_PROCESSOR = get_qwen2_vl_ocr_processor()
    QWEN2_VL_OCR_DEVICE = get_qwen2_vl_ocr_device()
    process_vision_info = get_process_vision_info()
    
    logger.debug(f"QWEN2_VL_OCR_MODEL: {QWEN2_VL_OCR_MODEL}")
    logger.debug(f"QWEN2_VL_OCR_PROCESSOR: {QWEN2_VL_OCR_PROCESSOR}")
    logger.debug(f"process_vision_info: {process_vision_info}")
    
    if not QWEN2_VL_OCR_MODEL or not QWEN2_VL_OCR_PROCESSOR or not process_vision_info:
        logger.error("Qwen2-VL-OCR model, processor, or process_vision_info not initialized. Cannot process PDF.")
        logger.error(f"Model is None: {QWEN2_VL_OCR_MODEL is None}")
        logger.error(f"Processor is None: {QWEN2_VL_OCR_PROCESSOR is None}")
        logger.error(f"process_vision_info is None: {process_vision_info is None}")
        return None, None

    full_document_text = ""
    page_level_metadata = []

    try:
        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        doc.close()
    except Exception as e:
        logger.error(f"Could not read PDF or get page count for {pdf_path}: {e}", exc_info=True)
        return None, None

    # Create a temporary directory for saving page images if needed by process_vision_info
    # or if passing PIL images directly is problematic.
    # For simplicity, we'll try to use PIL images directly.
    # If process_vision_info strictly needs paths, we'll need to save images temporarily.

    for i in range(num_pages):
        page_num_actual = i + 1
        logger.info(f"Processing page {page_num_actual}/{num_pages} of {pdf_path} with Qwen2-VL-OCR...")
        page_text = ""
        try:
            pil_image = render_pdf_page_to_image(pdf_path, i) # page index i
            if not pil_image:
                full_document_text += f"\n\n[OCR_ERROR_PAGE_{page_num_actual}: Failed to render page]\n\n"
                page_level_metadata.append({"page_number": page_num_actual, "status": "render_failed", "char_count": 0})
                continue

            # Construct messages for Qwen2-VL-OCR
            # The `process_vision_info` function from `qwen_vl_utils` in the example
            # expects image paths. If it can handle PIL.Image objects, that's better.
            # Let's assume for now we pass the PIL Image directly, and `process_vision_info`
            # or the processor handles it. If not, we'd save the PIL image to a temp file.
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "image", "image": pil_image}, # Pass PIL image first
                        {"type": "text", "text": QWEN2_VL_OCR_PROMPT}  # Text after image as per example
                    ],
                }
            ]

            # Process inputs
            # This follows the `prithivMLmods` example structure from `qwen2_vl.md`
            templated_text = QWEN2_VL_OCR_PROCESSOR.apply_chat_template(
                messages, tokenize=False, add_generation_prompt=True
            )
            
            # `process_vision_info` will extract the PIL.Image from `messages`
            # and prepare it.
            image_inputs, video_inputs = process_vision_info(messages)  # Don't pass processor as per example
            
            # Handle the case where video_inputs is an empty list
            if video_inputs is not None and len(video_inputs) == 0:
                video_inputs = None
                
            inputs = QWEN2_VL_OCR_PROCESSOR(
                text=[templated_text],
                images=image_inputs, # This should be the output from process_vision_info
                videos=video_inputs, # Should be None for images-only processing
                padding=True,
                return_tensors="pt",
            )
            # Move inputs to the correct device (where the model is)
            # QWEN2_VL_OCR_MODEL.device is the primary device if device_map was used.
            # Or use QWEN2_VL_OCR_DEVICE if model is on a single device.
            # For device_map="auto", inputs should be moved to QWEN2_VL_OCR_MODEL.device
            # However, the example shows inputs.to("cuda") if cuda is available.
            # Let's use QWEN2_VL_OCR_DEVICE which we set to 'cuda' or 'cpu'.
            inputs = inputs.to(QWEN2_VL_OCR_DEVICE)


            # Inference: Generation of the output
            generated_ids = QWEN2_VL_OCR_MODEL.generate(
                **inputs, 
                max_new_tokens=QWEN2_VL_OCR_MAX_NEW_TOKENS,
                temperature=0.3,  # Slightly higher than 0.1 for better results  
                do_sample=True,   # Enable sampling
                top_p=0.9,        # Nucleus sampling
                repetition_penalty=1.05  # Slight penalty to prevent repetition
            )
            
            # Decode output
            # The example trims input_ids from generated_ids.
            generated_ids_trimmed = [
                out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
            ]
            output_decoded = QWEN2_VL_OCR_PROCESSOR.batch_decode(
                generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
            )
            
            if output_decoded:
                page_text = output_decoded[0].strip()
            else:
                page_text = ""
            
            full_document_text += page_text + "\n\n"
            page_level_metadata.append({
                "page_number": page_num_actual,
                "method": "Qwen2-VL-OCR",
                "char_count": len(page_text)
            })
            logger.debug(f"Extracted {len(page_text)} chars from page {page_num_actual} of {pdf_path}.")

        except Exception as e:
            logger.error(f"Error processing page {page_num_actual} of {pdf_path} with Qwen2-VL-OCR: {e}", exc_info=True)
            full_document_text += f"\n\n[OCR_ERROR_PAGE_{page_num_actual}: {str(e)}]\n\n"
            page_level_metadata.append({
                "page_number": page_num_actual,
                "status": "ocr_failed",
                "error": str(e),
                "char_count": 0
            })
        finally:
            # Clean up PIL image if not needed anymore to save memory
            if 'pil_image' in locals() and pil_image:
                del pil_image
                if QWEN2_VL_OCR_DEVICE == 'cuda':
                    torch.cuda.empty_cache() # Clear cache if on GPU per page

    logger.info(f"Qwen2-VL-OCR processing completed for {pdf_path}. Total characters extracted: {len(full_document_text)}.")
    return full_document_text.strip(), page_level_metadata


# Comment out or remove extract_text_from_pdf_donut
# def extract_text_from_pdf_donut(pdf_path: str) -> tuple[str | None, list | None]:
#    ... (old Donut code) ...

# Keep other extraction functions like extract_text_from_pdf_olmocr, extract_text_from_docx, etc.


# Helper function for downloading from Supabase storage if path is a URL
def _download_supabase_file_to_temp(supabase_url: str) -> str | None:
    try:
        logger.info(f"Downloading from Supabase storage URL: {supabase_url}")
        response = requests.get(supabase_url, stream=True)
        response.raise_for_status()  # Raise an exception for HTTP errors
        
        # Create a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")  # Assuming PDF
        with open(temp_file.name, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        logger.info(f"Successfully downloaded to temporary file: {temp_file.name}")
        return temp_file.name
    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to download from Supabase URL {supabase_url}: {e}")
        return None
    except Exception as e:
        logger.error(f"Unexpected error downloading from Supabase URL {supabase_url}: {e}", exc_info=True)
        return None


def validate_pdf_for_processing(file_path: str) -> tuple[bool, str, dict]:
    """
    Pre-flight validation for PDF files before processing.
    
    Args:
        file_path: Path to the PDF file to validate
        
    Returns:
        tuple: (is_valid, error_message, metadata)
            - is_valid: True if PDF is valid for processing
            - error_message: Description of validation failure (empty if valid)
            - metadata: Dictionary with file details (size, pages, etc.)
    """
    metadata = {}
    
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}", metadata
            
        # Get file size
        file_size = os.path.getsize(file_path)
        metadata['file_size_bytes'] = file_size
        metadata['file_size_mb'] = file_size / (1024 * 1024)
        
        # Check file size limits (100MB max for Textract)
        if file_size > 100 * 1024 * 1024:
            return False, f"File too large: {metadata['file_size_mb']:.2f}MB (max 100MB)", metadata
            
        if file_size == 0:
            return False, "File is empty (0 bytes)", metadata
        
        # Validate PDF format and get page count
        try:
            with open(file_path, 'rb') as pdf_file:
                # Check PDF header
                header = pdf_file.read(5)
                if header != b'%PDF-':
                    return False, "Invalid PDF header - file may be corrupted", metadata
                
            # Get page count using PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(file_path)
                page_count = len(pdf_reader.pages)
                metadata['page_count'] = page_count
                
                # Check page limits (3000 pages max for Textract)
                if page_count > 3000:
                    return False, f"Too many pages: {page_count} (max 3000)", metadata
                    
                if page_count == 0:
                    return False, "PDF has no pages", metadata
                    
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    return False, "PDF is encrypted and cannot be processed", metadata
                    
            except Exception as e:
                # Try with PyMuPDF as fallback
                try:
                    pdf_doc = fitz.open(file_path)
                    page_count = len(pdf_doc)
                    metadata['page_count'] = page_count
                    pdf_doc.close()
                    
                    if page_count > 3000:
                        return False, f"Too many pages: {page_count} (max 3000)", metadata
                    if page_count == 0:
                        return False, "PDF has no pages", metadata
                        
                except Exception as e2:
                    return False, f"Could not read PDF structure: {str(e)}", metadata
        
        except Exception as e:
            return False, f"Invalid PDF file: {str(e)}", metadata
            
        # All validations passed
        metadata['validation_status'] = 'passed'
        return True, "", metadata
        
    except Exception as e:
        logger.error(f"Error during PDF validation: {e}", exc_info=True)
        return False, f"Validation error: {str(e)}", metadata


def extract_text_from_pdf_textract(db_manager: SupabaseManager,  # Added db_manager
                                   source_doc_sql_id: int,  # Added source_doc_sql_id
                                   pdf_path_or_s3_uri: str,
                                   document_uuid_from_db: str  # Expect this from DB
                                   ) -> tuple[str | None, list | None]:
    
    logger.info(f"Starting PDF text extraction with AWS Textract for source_doc_id: {source_doc_sql_id}, doc_uuid: {document_uuid_from_db}, path: {pdf_path_or_s3_uri}")
    
    # Check cache first
    try:
        textract_processor = TextractProcessor(db_manager=db_manager)
        cached_result = textract_processor.get_cached_ocr_result(document_uuid_from_db)
        if cached_result:
            text, metadata = cached_result
            logger.info(f"Using cached OCR result for document {document_uuid_from_db}")
            return text, metadata.get('page_metadata', [])
    except Exception as e:
        logger.debug(f"Error checking OCR cache: {e}")
    
    s3_manager = S3StorageManager()
    # Pass db_manager to TextractProcessor for DB updates
    textract_processor = TextractProcessor(db_manager=db_manager)

    s3_bucket_name = S3_PRIMARY_DOCUMENT_BUCKET
    s3_object_key: str | None = None
    local_temp_pdf_to_clean: str | None = None

    try:
        if pdf_path_or_s3_uri.startswith('s3://'):
            parts = pdf_path_or_s3_uri.replace('s3://', '').split('/', 1)
            if len(parts) == 2:
                s3_bucket_name = parts[0]  # Could be different from S3_PRIMARY_DOCUMENT_BUCKET if already processed
                s3_object_key = parts[1]
                logger.info(f"Processing existing S3 object: s3://{s3_bucket_name}/{s3_object_key}")
            else:  # Should not happen if queue processor formats correctly
                logger.error(f"Invalid S3 URI format: {pdf_path_or_s3_uri}")
                db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_S3_FORMAT", 'failed')
                return None, [{"status": "error", "error_message": f"Invalid S3 URI: {pdf_path_or_s3_uri}"}]
        
        elif pdf_path_or_s3_uri.startswith('supabase://'):
            # Handle Supabase Storage URLs - convert to HTTP URL
            logger.info(f"Supabase Storage URL detected: {pdf_path_or_s3_uri}")
            # Extract bucket and path from supabase://bucket/path format
            parts = pdf_path_or_s3_uri.replace('supabase://', '').split('/', 1)
            if len(parts) == 2:
                bucket = parts[0]
                path = parts[1]
                # Construct the HTTP URL for Supabase Storage
                supabase_url = f"{db_manager.client.supabase_url}/storage/v1/object/public/{bucket}/{path}"
                logger.info(f"Converted to HTTP URL: {supabase_url}")
                local_temp_pdf_to_clean = _download_supabase_file_to_temp(supabase_url)
                if not local_temp_pdf_to_clean:
                    db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_SUPABASE_DL_FAIL", 'failed')
                    return None, [{"status": "error", "error_message": f"Failed to download from Supabase Storage: {pdf_path_or_s3_uri}"}]
            else:
                logger.error(f"Invalid Supabase Storage URL format: {pdf_path_or_s3_uri}")
                db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_SUPABASE_FORMAT", 'failed')
                return None, [{"status": "error", "error_message": f"Invalid Supabase Storage URL: {pdf_path_or_s3_uri}"}]
                
        elif pdf_path_or_s3_uri.startswith('http://') or pdf_path_or_s3_uri.startswith('https://'):
            # Assume it's a Supabase Storage URL or other accessible HTTP URL
            logger.info(f"HTTP(S) URL detected: {pdf_path_or_s3_uri}. Downloading temporarily.")
            local_temp_pdf_to_clean = _download_supabase_file_to_temp(pdf_path_or_s3_uri)
            if not local_temp_pdf_to_clean:
                db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_HTTP_DL_FAIL", 'failed')
                return None, [{"status": "error", "error_message": f"Failed to download from URL: {pdf_path_or_s3_uri}"}]
            
            # Now upload this temporary local file to S3 using the document_uuid_from_db
            # The original_filename for S3 metadata can be derived or fixed.
            # If source_documents table has original_file_name, fetch and use it.
            source_doc_details = db_manager.get_document_by_id(source_doc_sql_id)
            original_filename_for_s3 = source_doc_details.get('original_file_name', f"{document_uuid_from_db}.pdf") if source_doc_details else f"{document_uuid_from_db}.pdf"

            upload_info = s3_manager.upload_document_with_uuid_naming(
                local_file_path=local_temp_pdf_to_clean,
                document_uuid=document_uuid_from_db,
                original_filename=original_filename_for_s3
            )
            s3_object_key = upload_info['s3_key']
            s3_bucket_name = upload_info['s3_bucket']  # Should be S3_PRIMARY_DOCUMENT_BUCKET
            logger.info(f"Uploaded downloaded file to s3://{s3_bucket_name}/{s3_object_key}")
            # Update source_documents with the new S3 key/bucket if this is a permanent migration
            db_manager.client.table('source_documents').update({
                's3_key': s3_object_key,
                's3_bucket': s3_bucket_name,
                's3_region': AWS_DEFAULT_REGION  # from config
            }).eq('id', source_doc_sql_id).execute()
            logger.info(f"Updated source_document {source_doc_sql_id} with new S3 location.")

        elif os.path.exists(pdf_path_or_s3_uri):
            logger.info(f"Local file detected: {pdf_path_or_s3_uri}. Validating before S3 upload.")
            original_filename = os.path.basename(pdf_path_or_s3_uri)
            
            # Pre-flight validation for PDFs
            if pdf_path_or_s3_uri.lower().endswith('.pdf'):
                is_valid, validation_error, validation_metadata = validate_pdf_for_processing(pdf_path_or_s3_uri)
                
                if not is_valid:
                    logger.error(f"PDF validation failed: {validation_error}")
                    
                    # Store validation error in ocr_metadata_json
                    error_metadata = [{
                        "status": "error",
                        "stage": "pdf_validation",
                        "error_message": validation_error,
                        "validation_metadata": validation_metadata,
                        "file_path": pdf_path_or_s3_uri,
                        "document_uuid": document_uuid_from_db,
                        "timestamp": datetime.now(timezone.utc).isoformat()
                    }]
                    
                    # Update database with validation error
                    db_manager.client.table('source_documents').update({
                        'textract_job_id': 'N/A_PDF_VALIDATION_FAILED',
                        'textract_job_status': 'failed',
                        'ocr_metadata_json': error_metadata,
                        'error_message': f"PDF Validation Failed: {validation_error}"
                    }).eq('id', source_doc_sql_id).execute()
                    
                    return None, error_metadata
                
                logger.info(f"✅ PDF validation passed: {validation_metadata}")
            
            # Enhanced error handling for S3 upload
            try:
                upload_info = s3_manager.upload_document_with_uuid_naming(
                    local_file_path=pdf_path_or_s3_uri,
                    document_uuid=document_uuid_from_db,
                    original_filename=original_filename
                )
                s3_object_key = upload_info['s3_key']
                s3_bucket_name = upload_info['s3_bucket']
                logger.info(f"✅ Successfully uploaded local file to s3://{s3_bucket_name}/{s3_object_key}")
                
            except ClientError as e:
                error_code = e.response.get('Error', {}).get('Code', 'Unknown')
                error_message = e.response.get('Error', {}).get('Message', str(e))
                logger.error(f"AWS S3 upload failed - Code: {error_code}, Message: {error_message}")
                
                # Store detailed error in ocr_metadata_json
                error_metadata = [{
                    "status": "error",
                    "stage": "s3_upload",
                    "error_code": error_code,
                    "error_message": error_message,
                    "file_path": pdf_path_or_s3_uri,
                    "document_uuid": document_uuid_from_db,
                    "timestamp": datetime.now(timezone.utc).isoformat()
                }]
                
                # Update database with detailed error
                db_manager.client.table('source_documents').update({
                    'textract_job_id': 'N/A_S3_UPLOAD_FAILED',
                    'textract_job_status': 'failed',
                    'ocr_metadata_json': error_metadata,
                    'error_message': f"S3 Upload Error ({error_code}): {error_message}"
                }).eq('id', source_doc_sql_id).execute()
                
                return None, error_metadata
                
            except Exception as e:
                logger.error(f"Unexpected error during S3 upload: {type(e).__name__}: {str(e)}")
                
                # Store detailed error in ocr_metadata_json
                error_metadata = [{
                    "status": "error", 
                    "stage": "s3_upload",
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "file_path": pdf_path_or_s3_uri,
                    "document_uuid": document_uuid_from_db,
                    "timestamp": datetime.now(timezone.utc).isoformat()
                }]
                
                # Update database with detailed error
                db_manager.client.table('source_documents').update({
                    'textract_job_id': 'N/A_S3_UPLOAD_EXCEPTION',
                    'textract_job_status': 'failed',
                    'ocr_metadata_json': error_metadata,
                    'error_message': f"S3 Upload Exception: {type(e).__name__}: {str(e)}"
                }).eq('id', source_doc_sql_id).execute()
                
                return None, error_metadata
            # Update source_documents with S3 key/bucket
            db_manager.client.table('source_documents').update({
                's3_key': s3_object_key,
                's3_bucket': s3_bucket_name,
                's3_region': AWS_DEFAULT_REGION
            }).eq('id', source_doc_sql_id).execute()
            logger.info(f"Updated source_document {source_doc_sql_id} with S3 location from local upload.")
        else:
            logger.error(f"File not found at path: {pdf_path_or_s3_uri}")
            db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_PATH_NOT_FOUND", 'failed')
            return None, [{"status": "error", "error_message": f"File not found: {pdf_path_or_s3_uri}"}]

        if not s3_object_key:  # Should have been set by one of the branches above
            logger.error(f"S3 object key could not be determined for {pdf_path_or_s3_uri}")
            db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_S3_KEY_MISSING", 'failed')
            return None, [{"status": "error", "error_message": "S3 object key missing"}]

        # Check S3 object existence
        if not s3_manager.check_s3_object_exists(s3_key=s3_object_key, s3_bucket=s3_bucket_name):
            logger.error(f"S3 object s3://{s3_bucket_name}/{s3_object_key} not found or accessible.")
            db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_S3_NOT_FOUND", 'failed')
            return None, [{"status": "error", "error_message": f"S3 object s3://{s3_bucket_name}/{s3_object_key} not found"}]

        start_time = time.time()
        extracted_text = None
        textract_doc_api_metadata = None  # From Textract GetDocumentTextDetection response
        page_level_metadata_for_db = []  # This is the ocr_metadata_json for source_documents

        if TEXTRACT_USE_ASYNC_FOR_PDF:
            logger.info(f"Using ASYNC Textract processing for s3://{s3_bucket_name}/{s3_object_key}")
            # Pass source_doc_sql_id and document_uuid_from_db to start_document_text_detection
            job_id = textract_processor.start_document_text_detection(
                s3_bucket=s3_bucket_name,
                s3_key=s3_object_key,
                source_doc_id=source_doc_sql_id,
                document_uuid_from_db=document_uuid_from_db
                # client_request_token and job_tag are handled inside start_document_text_detection
            )

            if not job_id:  # Failure to start job (already logged and DB updated by start_document_text_detection)
                return None, [{"status": "error", "error_message": f"Failed to start Textract job for s3://{s3_bucket_name}/{s3_object_key}"}]

            # Pass source_doc_sql_id to get_text_detection_results for status updates
            blocks, textract_doc_api_metadata = textract_processor.get_text_detection_results(job_id, source_doc_sql_id)
            
            if blocks:
                extracted_text = textract_processor.process_textract_blocks_to_text(blocks, textract_doc_api_metadata)
            else:  # Job might have succeeded but returned no blocks, or failed. get_text_detection_results handles DB updates.
                logger.error(f"No blocks returned from Textract job {job_id}. Check textract_jobs table for details.")
        else:  # Synchronous
            logger.error("Synchronous PDF processing via DetectDocumentText is not fully supported for multi-page. Configure TEXTRACT_USE_ASYNC_FOR_PDF=true.")
            db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_SYNC_UNSUPPORTED", 'failed')
            return None, [{"status": "error", "error_message": "Sync PDF Textract not supported"}]

        processing_time_total = time.time() - start_time
        
        # Create page-level metadata for ocr_metadata_json in source_documents
        if extracted_text and textract_doc_api_metadata:
            num_pages = textract_doc_api_metadata.get('Pages', 1)
            page_level_metadata_for_db = [{
                "page_number": i + 1,
                "method": "AWSTextract",
                "confidence_threshold": TEXTRACT_CONFIDENCE_THRESHOLD,
                "processing_time_seconds": processing_time_total / num_pages,  # Averaged per page
                "s3_key": s3_object_key,
                "s3_bucket": s3_bucket_name
            } for i in range(num_pages)]
            
        # The final update of source_documents with extracted_text is handled by the caller or within TextractProcessor's methods.
        # Specifically, get_text_detection_results calls update_source_document_with_textract_outcome with various fields.
        # However, the `raw_extracted_text` field update needs to be here or done by the caller.
        if extracted_text:
            db_manager.update_source_document_with_textract_outcome(
                source_doc_sql_id=source_doc_sql_id,
                textract_job_id=job_id if 'job_id' in locals() else "SYNC_NOT_IMPL",
                textract_job_status='succeeded',
                raw_text=extracted_text,
                ocr_metadata=page_level_metadata_for_db  # This is what goes into ocr_metadata_json
            )
            
            # Cache the OCR result
            try:
                textract_processor._cache_ocr_result(
                    document_uuid_from_db, 
                    extracted_text, 
                    {'page_metadata': page_level_metadata_for_db}
                )
            except Exception as e:
                logger.debug(f"Error caching OCR result: {e}")
        
        return extracted_text, page_level_metadata_for_db
    
    except Exception as e:
        logger.error(f"Unexpected error in extract_text_from_pdf_textract: {e}", exc_info=True)
        db_manager.update_source_document_with_textract_outcome(source_doc_sql_id, "N/A_EXCEPTION", 'failed')
        return None, [{"status": "error", "error_message": str(e)}]
    finally:
        # Clean up temporary downloaded file if any
        if local_temp_pdf_to_clean and os.path.exists(local_temp_pdf_to_clean):
            try:
                os.unlink(local_temp_pdf_to_clean)
                logger.info(f"Cleaned up temporary file: {local_temp_pdf_to_clean}")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file {local_temp_pdf_to_clean}: {e}")

def extract_text_from_pdf_olmocr(pdf_path: str) -> tuple[str | None, list | None]:
    """Extracts text from all pages of a PDF using olmOCR. This is a placeholder function."""
    # OLMOCR is not implemented in this version
    logger.error("olmOCR model not implemented in this version. Cannot process PDF.")
    return None, None


def extract_text_from_docx(docx_path: str) -> str | None:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(docx_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
        return None


def extract_text_from_docx_s3_aware(file_path_or_s3_uri: str, s3_manager=None) -> tuple[str | None, list | None]:
    """
    Extract text from DOCX that may be local or in S3.
    Returns (text, metadata) tuple like other extraction functions.
    """
    import tempfile
    import os
    from docx import Document as DocxDocument
    
    local_temp_file = None
    
    try:
        # Handle S3 URLs
        if file_path_or_s3_uri.startswith('s3://'):
            if not s3_manager:
                from scripts.s3_storage import S3StorageManager
                s3_manager = S3StorageManager()
            
            # Parse S3 URL
            parts = file_path_or_s3_uri.replace('s3://', '').split('/', 1)
            if len(parts) != 2:
                logger.error(f"Invalid S3 URI format: {file_path_or_s3_uri}")
                return None, [{"status": "error", "error_message": "Invalid S3 URI"}]
            
            bucket_name, s3_key = parts
            
            # Download to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                local_temp_file = tmp_file.name
                s3_manager.s3_client.download_file(bucket_name, s3_key, local_temp_file)
                logger.info(f"Downloaded DOCX from S3 to {local_temp_file}")
            
            docx_path = local_temp_file
        else:
            # Local file - strip any suffix
            if '_' in file_path_or_s3_uri and len(file_path_or_s3_uri.split('_')[-1]) == 8:
                docx_path = '_'.join(file_path_or_s3_uri.split('_')[:-1])
            else:
                docx_path = file_path_or_s3_uri
        
        # Extract text
        doc = DocxDocument(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()
        
        # Get approximate page count (Word doesn't store exact page count in DOCX)
        # Using section breaks as approximation
        try:
            page_count = len(doc.element.xpath('//w:sectPr'))
            if page_count == 0:
                page_count = 1  # At least one page
        except:
            page_count = 1
        
        metadata = [{
            "method": "docx_parser",
            "pages": page_count,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }]
        
        return text, metadata
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path_or_s3_uri}: {e}")
        return None, [{"status": "error", "error_message": str(e)}]
    
    finally:
        # Clean up temp file
        if local_temp_file and os.path.exists(local_temp_file):
            try:
                os.unlink(local_temp_file)
                logger.debug(f"Cleaned up temp file {local_temp_file}")
            except Exception as e:
                logger.warning(f"Failed to clean up temp file {local_temp_file}: {e}")

def extract_text_from_txt(txt_path: str) -> str | None:
    try:
        with open(txt_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT {txt_path}: {e}")
        return None

def extract_text_from_eml(eml_path: str) -> str | None:
    try:
        import email
        from email import policy
        import re # For basic HTML stripping if needed

        with open(eml_path, 'rb') as f:
            msg = email.message_from_bytes(f.read(), policy=policy.default)

        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                cdispo = str(part.get('Content-Disposition'))
                if ctype == 'text/plain' and 'attachment' not in cdispo:
                    payload = part.get_payload(decode=True)
                    charset = part.get_content_charset() or 'utf-8'
                    body = payload.decode(charset, errors='replace')
                    break
            if not body: # Fallback to HTML part if no plain text
                for part in msg.walk():
                    ctype = part.get_content_type()
                    cdispo = str(part.get('Content-Disposition'))
                    if 'text/html' in ctype and 'attachment' not in cdispo:
                        html_payload = part.get_payload(decode=True)
                        html_charset = part.get_content_charset() or 'utf-8'
                        html_body = html_payload.decode(html_charset, errors='replace')
                        # Super basic HTML strip - consider BeautifulSoup4 for robustness
                        body = re.sub(r'<style.*?</style>', '', html_body, flags=re.DOTALL | re.IGNORECASE)
                        body = re.sub(r'<script.*?</script>', '', body, flags=re.DOTALL | re.IGNORECASE)
                        body = re.sub(r'<[^>]+>', ' ', body) # Replace tags with space
                        body = ' '.join(body.split()) # Normalize whitespace
                        break
        else:
            payload = msg.get_payload(decode=True)
            charset = msg.get_content_charset() or 'utf-8'
            body = payload.decode(charset, errors='replace')

        headers_text = f"From: {msg.get('From', '')}\nTo: {msg.get('To', '')}\nCC: {msg.get('CC','')}\nBCC: {msg.get('BCC','')}\nSubject: {msg.get('Subject', '')}\nDate: {msg.get('Date', '')}\n\n"
        return headers_text + body.strip()
    except Exception as e:
        logger.error(f"Error extracting text from EML {eml_path}: {e}")
        return None


def transcribe_audio_whisper(audio_path: str) -> str | None:
    """
    Transcribe audio using Whisper model.
    Stage 1: Uses OpenAI Whisper API
    Stage 2+: Uses local Whisper model
    """
    # Stage 1: Use OpenAI Whisper API
    if DEPLOYMENT_STAGE == STAGE_CLOUD_ONLY or USE_OPENAI_FOR_AUDIO_TRANSCRIPTION or not should_load_local_models():
        return transcribe_audio_openai_whisper(audio_path)
    
    # Stage 2+: Use local Whisper model
    WHISPER_MODEL = get_whisper_model()
    if not WHISPER_MODEL:
        logger.error("Whisper model not initialized. Cannot transcribe audio.")
        return None
    try:
        # Ensure fp16 is False if on CPU
        use_fp16 = False
        if hasattr(WHISPER_MODEL, 'device') and WHISPER_MODEL.device.type == 'cuda':
            use_fp16 = True

        result = WHISPER_MODEL.transcribe(audio_path, fp16=use_fp16)
        return result["text"].strip()
    except Exception as e:
        logger.error(f"Error transcribing audio {audio_path} with local Whisper: {e}")
        return None


def transcribe_audio_openai_whisper(audio_path: str) -> str | None:
    """
    Transcribe audio using OpenAI Whisper API for Stage 1 deployment.
    
    Args:
        audio_path: Path to the audio file to transcribe
        
    Returns:
        str: Transcribed text or None if failed
    """
    if not OpenAI:
        logger.error("OpenAI library not available. Cannot transcribe audio with OpenAI Whisper.")
        return None
        
    if not OPENAI_API_KEY:
        logger.error("OpenAI API key not configured. Cannot transcribe audio.")
        return None
    
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        # Check if file exists and get size
        if not os.path.exists(audio_path):
            logger.error(f"Audio file not found: {audio_path}")
            return None
            
        file_size = os.path.getsize(audio_path)
        logger.info(f"Transcribing audio file: {audio_path} (size: {file_size} bytes)")
        
        # OpenAI Whisper API has a 25MB file size limit
        if file_size > 25 * 1024 * 1024:  # 25MB
            logger.error(f"Audio file too large for OpenAI Whisper API: {file_size / (1024*1024):.2f}MB (limit: 25MB)")
            return None
        
        # Open and transcribe the audio file
        with open(audio_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="text"
            )
            
        transcribed_text = transcript.strip() if transcript else ""
        logger.info(f"OpenAI Whisper transcription completed. Text length: {len(transcribed_text)} characters")
        
        # Log a sample of the transcribed text for debugging
        if transcribed_text:
            sample = transcribed_text[:200] + "..." if len(transcribed_text) > 200 else transcribed_text
            logger.debug(f"Transcription sample: {sample}")
        
        return transcribed_text
        
    except Exception as e:
        logger.error(f"Error transcribing audio {audio_path} with OpenAI Whisper: {e}", exc_info=True)
        return None